import os
import logging
import subprocess

from google.cloud import storage
from google.cloud import documentai_v1 as documentai
from google.oauth2 import service_account

from dotenv import load_dotenv
from docx import Document

# Configure logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s [%(levelname)s] %(message)s', datefmt='%Y-%m-%d %H:%M:%S')

# Load environment variables from .env file
logging.debug("Loading environment variables from .env file")
load_dotenv()

# Load credentials from the JSON key file
SERVICE_ACCOUNT_FILE = os.getenv('GOOGLE_APPLICATION_CREDENTIALS')
SCOPES = ['https://www.googleapis.com/auth/cloud-platform']  # Adjust scopes if needed

logging.debug(f"Service Account File: {SERVICE_ACCOUNT_FILE}")
logging.debug(f"Scopes: {SCOPES}")

try:
    credentials = service_account.Credentials.from_service_account_file(SERVICE_ACCOUNT_FILE, scopes=SCOPES)
    logging.debug("Service account credentials loaded successfully")
except Exception as e:
    logging.error(f"Error loading service account credentials: {e}")
    raise

# Cloud Storage settings
BUCKET_NAME = os.getenv('CLOUD_STORAGE_BUCKET')
ATTACHMENT_FOLDER = "attachments"  # Make sure this matches the first script

logging.debug(f"Cloud Storage Bucket: {BUCKET_NAME}")

def sanitize_text(text):
    # Replace Unicode characters that LaTeX cannot parse with '.'
    return ''.join(char if ord(char) < 128 else '.' for char in text)

def sanitize_docx(input_path, sanitized_path):
    doc = Document(input_path)
    sanitized_doc = Document()

    for para in doc.paragraphs:
        sanitized_paragraph = sanitize_text(para.text)
        sanitized_doc.add_paragraph(sanitized_paragraph)

    sanitized_doc.save(sanitized_path)

def process_attachments(event, context):
    """Triggered by a change to a Cloud Storage bucket.
    Args:
         event (dict): Event payload.
         context (google.cloud.functions.Context): Metadata for the event.
    """
    logging.debug(f"Event: {event}")
    logging.debug(f"Context: {context}")

    file_name = event['name']
    logging.debug(f"Processing file: {file_name}")

    # Check if file is in the attachments folder and has a supported extension
    if not file_name.startswith(ATTACHMENT_FOLDER + "/") or not any(file_name.endswith(ext) for ext in ['.doc', '.docx', '.xls']):
        logging.info(f"Skipping file: {file_name} (not in attachments folder or unsupported extension)")
        return

    # Initialize clients
    storage_client = storage.Client(credentials=credentials)
    documentai_client = documentai.DocumentProcessorServiceClient(credentials=credentials)

    # Get file from bucket
    bucket = storage_client.bucket(BUCKET_NAME)
    blob = bucket.blob(file_name)

    # Download file to temporary location
    temp_file_path = f"/tmp/{file_name.split('/')[-1]}"
    blob.download_to_filename(temp_file_path)

    if file_name.endswith(".docx"):
        # Sanitize and convert .docx to PDF using Pandoc
        sanitized_path = os.path.join("/tmp", "sanitized_" + file_name.split('/')[-1])
        output_path = os.path.join("/tmp", os.path.splitext(file_name.split('/')[-1])[0] + ".pdf")

        logging.debug(f"Sanitizing {temp_file_path}...")
        try:
            sanitize_docx(temp_file_path, sanitized_path)

            logging.debug(f"Converting {sanitized_path} to PDF...")
            subprocess.run(['pandoc', sanitized_path, '-o', output_path], check=True)
            logging.debug(f"Converted: {output_path}")

            # Upload the PDF to Cloud Storage
            pdf_file_name = f"{ATTACHMENT_FOLDER}/{os.path.splitext(file_name.split('/')[-1])[0]}.pdf"
            pdf_blob = bucket.blob(pdf_file_name)
            with open(output_path, "rb") as f:
                pdf_blob.upload_from_file(f, content_type="application/pdf")

        except subprocess.CalledProcessError as e:
            logging.error(f"Error converting {temp_file_path}: {str(e)}")

        finally:
            # Clean up temporary files
            os.remove(sanitized_path)
            os.remove(output_path)

    else:  # .doc or .xls
        # Process with Document AI
        with open(temp_file_path, "rb") as f:
            raw_document = documentai.types.RawDocument(content=f.read(), mime_type="application/octet-stream")

        name = f"projects/{os.getenv('PROJECT_ID')}/locations/us/processors/my-doc-processor"
        logging.debug(f"Using Document AI processor: {name}")

        result = documentai_client.process_document(name=name, raw_document=raw_document)
        pdf_bytes = result.document.content

        # Upload PDF back to bucket
        pdf_file_name = f"{ATTACHMENT_FOLDER}/{os.path.splitext(file_name.split('/')[-1])[0]}.pdf"
        pdf_blob = bucket.blob(pdf_file_name)
        pdf_blob.upload_from_string(pdf_bytes, content_type="application/pdf")

    # Delete the original file
    blob.delete()

    logging.info(f"Processed {file_name} into {pdf_file_name}")